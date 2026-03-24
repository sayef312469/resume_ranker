# ============================================================
# cv_parser.py
# Extracts text from PDF, DOCX, and TXT files
# ============================================================

import io
import re
import os
import warnings
warnings.filterwarnings("ignore")

from fastapi import UploadFile


# ============================================================
# INSTALL CHECK — run these if not installed:
# pip install PyPDF2 python-docx
# ============================================================


# ============================================================
# MAIN EXTRACTION FUNCTION
# Detects file type and routes to correct parser
# ============================================================

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Main entry point — detects file type and extracts text.
    Supports: PDF, DOCX, DOC, TXT
    Returns: raw text string
    """
    filename  = file.filename.lower()
    contents  = await file.read()

    # Reset file pointer for potential re-reads
    await file.seek(0)

    if filename.endswith('.pdf'):
        text = extract_from_pdf(contents)

    elif filename.endswith('.docx'):
        text = extract_from_docx(contents)

    elif filename.endswith('.txt'):
        text = extract_from_txt(contents)

    elif filename.endswith('.doc'):
        # .doc is old Word format — harder to parse
        # Fall back to raw text extraction
        text = extract_from_txt(contents)

    else:
        raise ValueError(
            f"Unsupported file type: {filename}. "
            f"Please upload PDF, DOCX, or TXT."
        )

    # Clean extracted text
    text = clean_extracted_text(text)

    return text


# ============================================================
# PDF PARSER
# ============================================================

def extract_from_pdf(contents: bytes) -> str:
    """
    Extracts text from PDF file bytes.
    Tries PyPDF2 first then falls back to pdfminer.
    """
    text = ""

    # Method 1 — PyPDF2 (fast, works for most PDFs)
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(
            io.BytesIO(contents)
        )
        pages_text = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)

        text = "\n".join(pages_text)

        if len(text.strip()) > 50:
            return text

    except Exception as e:
        print(f"   PyPDF2 failed: {e} — trying pdfminer...")

    # Method 2 — pdfminer (slower but more accurate)
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        import io as _io

        output_string = _io.StringIO()
        extract_text_to_fp(
            io.BytesIO(contents),
            output_string,
            laparams = LAParams()
        )
        text = output_string.getvalue()

        if len(text.strip()) > 50:
            return text

    except ImportError:
        print("   pdfminer not installed — "
              "pip install pdfminer.six")
    except Exception as e:
        print(f"   pdfminer failed: {e}")

    # If both fail return what we have
    return text if text else ""


# ============================================================
# DOCX PARSER
# ============================================================

def extract_from_docx(contents: bytes) -> str:
    """
    Extracts text from DOCX file bytes.
    Extracts paragraphs + tables.
    """
    try:
        import docx

        doc        = docx.Document(io.BytesIO(contents))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")


# ============================================================
# TXT PARSER
# ============================================================

def extract_from_txt(contents: bytes) -> str:
    """
    Extracts text from plain text file bytes.
    Tries multiple encodings.
    """
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']

    for encoding in encodings:
        try:
            return contents.decode(encoding)
        except UnicodeDecodeError:
            continue

    # Last resort — ignore errors
    return contents.decode('utf-8', errors='ignore')


# ============================================================
# TEXT CLEANING
# Cleans extracted raw text
# ============================================================

def clean_extracted_text(text: str) -> str:
    """
    Cleans raw extracted text:
    - Removes excessive whitespace
    - Removes special characters
    - Normalizes line breaks
    """
    if not text:
        return ""

    # Normalize line breaks
    text = text.replace('\r\n', '\n')
    text = text.replace('\r', '\n')

    # Remove excessive blank lines
    lines      = text.split('\n')
    lines      = [line.strip() for line in lines]
    lines      = [
        line for line in lines
        if line  # remove empty lines
    ]
    text       = '\n'.join(lines)

    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)

    # Remove excessive spaces
    text = re.sub(r' {3,}', ' ', text)

    # Remove very short lines (likely noise)
    lines = text.split('\n')
    lines = [
        line for line in lines
        if len(line.strip()) > 2
    ]
    text  = '\n'.join(lines)

    return text.strip()


# ============================================================
# VALIDATE EXTRACTED TEXT
# Makes sure extracted text is usable
# ============================================================

def validate_cv_text(text: str) -> dict:
    """
    Validates that extracted text looks like a CV.
    Returns validation result with details.
    """
    if not text:
        return {
            "valid"  : False,
            "reason" : "No text extracted from file"
        }

    word_count = len(text.split())
    char_count = len(text)

    if char_count < 100:
        return {
            "valid"  : False,
            "reason" : f"Text too short ({char_count} chars). "
                       f"File may be image-based or corrupted."
        }

    if word_count < 30:
        return {
            "valid"  : False,
            "reason" : f"Too few words ({word_count}). "
                       f"CV may not have been extracted properly."
        }

    # Check for common CV keywords
    cv_keywords = [
        'experience', 'education', 'skills',
        'work', 'university', 'bachelor',
        'master', 'project', 'internship',
        'resume', 'cv', 'objective', 'summary'
    ]
    text_lower  = text.lower()
    found_kws   = [
        kw for kw in cv_keywords
        if kw in text_lower
    ]

    if len(found_kws) < 2:
        return {
            "valid"  : True,  # still valid, just warn
            "warning": "Text may not be a CV — "
                       "few CV keywords detected",
            "reason" : None
        }

    return {
        "valid"     : True,
        "word_count": word_count,
        "char_count": char_count,
        "cv_keywords_found": found_kws
    }


# ============================================================
# HELPER — PREVIEW TEXT
# Returns first N words for preview purposes
# ============================================================

def preview_text(text: str, n_words: int = 50) -> str:
    """Returns first N words of text for preview"""
    words = text.split()[:n_words]
    return " ".join(words) + "..." if len(
        text.split()
    ) > n_words else text